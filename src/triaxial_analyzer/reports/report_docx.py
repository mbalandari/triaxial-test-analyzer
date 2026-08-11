"""
Generate DOCX reports.
"""

from docx import Document


def generate_docx_report(results, envelope, plot_path, output_path):
    """
    Create a DOCX report including Mohr plot.

    Parameters:
        results: List of AnalysisResult objects.
        envelope: Envelope dict.
        plot_path: Path to saved plot image.
        output_path: DOCX file path.

    Returns:
        None
    """
    doc = Document()

    doc.add_heading("Triaxial Test Analysis Report", level=1)

    for r in results:
        doc.add_heading("Specimen", level=2)
        doc.add_paragraph(f"Peak Stress: {r.peak_stress:.3f} MPa")
        doc.add_paragraph(f"Peak Strain: {r.peak_strain:.5f}")
        doc.add_paragraph(f"Sigma1: {r.sigma1:.3f} MPa")
        doc.add_paragraph(f"Sigma3: {r.sigma3:.3f} MPa")

    doc.add_heading("Mohr–Coulomb Envelope", level=2)
    doc.add_paragraph(f"Cohesion: {envelope['cohesion']:.3f} MPa")
    doc.add_paragraph(f"Friction Angle: {envelope['phi_deg']:.2f}°")

    doc.add_picture(plot_path)

    doc.save(output_path)
